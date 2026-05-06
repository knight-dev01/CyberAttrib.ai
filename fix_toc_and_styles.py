
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def add_toc(paragraph):
    # Inserts a standard Word TOC field that Word will recognize
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._element.append(fldChar)
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    run._element.append(instrText)
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    run._element.append(fldChar2)
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._element.append(fldChar3)

def add_page_number(doc):
    # Add page numbers to footer
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run()
        
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run._element.append(fldChar1)
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE'
        run._element.append(instrText)
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._element.append(fldChar2)

def standardise_and_fix_toc():
    doc_path = r'c:\Users\Adegoke\Desktop\New folder\docs\ADIPHAS\ADIPHAS_Final_Submission_Perfected.docx'
    output_path = r'c:\Users\Adegoke\Desktop\New folder\docs\ADIPHAS\ADIPHAS_Submission_Final_Fixed.docx'
    
    doc = docx.Document(doc_path)
    
    # 1. FIX STYLES FOR TOC (Critical)
    for p in doc.paragraphs:
        # Chapter Headings
        if "CHAPTER" in p.text.upper() and len(p.text) < 100:
            p.style = 'Heading 1'
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(16)
                run.font.name = 'Times New Roman'
        
        # Section Headings (e.g., 1.1, 2.3.1)
        elif re.match(r"^\d+\.\d+", p.text.strip()):
            p.style = 'Heading 2'
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(14)
                run.font.name = 'Times New Roman'

        # Abstract, Dedication, Acknowledgment, References
        elif p.text.strip().upper() in ["ABSTRACT", "DEDICATION", "ACKNOWLEDGEMENT", "REFERENCES", "TABLE OF CONTENTS"]:
            p.style = 'Heading 1'
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(16)
                run.font.name = 'Times New Roman'

    # 2. RE-INSERT TOC (Ensure it's fresh)
    toc_found = False
    for i, p in enumerate(doc.paragraphs):
        if "TABLE OF CONTENTS" in p.text.upper():
            # Clear following paragraph (where old TOC was)
            if i + 1 < len(doc.paragraphs):
                doc.paragraphs[i+1].text = ""
                add_toc(doc.paragraphs[i+1])
            toc_found = True
            break
            
    # 3. PAGE NUMBERS
    add_page_number(doc)

    # 4. MARGINS (Standard 1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    doc.save(output_path)
    print(f"Success: Fixed TOC and standardized document saved to {output_path}")

standardise_and_fix_toc()
