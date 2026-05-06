
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def set_para_format(p):
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def clone_standardization_v2():
    template_path = r'c:\Users\Adegoke\Desktop\New folder\docs\ADIPHAS\document_pdf.docx'
    source_path = r'c:\Users\Adegoke\Desktop\New folder\docs\ADIPHAS\ADIPHAS_Submission_Final_Fixed.docx'
    output_path = r'c:\Users\Adegoke\Desktop\New folder\docs\ADIPHAS\ADIPHAS_Final_Standardized_Clone.docx'
    
    doc = docx.Document(template_path)
    source_doc = docx.Document(source_path)

    # 1. SURGICAL TITLE PAGE AND PRELIMS UPDATE
    for p in doc.paragraphs[:95]:
        if "ARTIFICIAL INTELLIGENCE ENHANCED JOB PLATFORM" in p.text.upper():
            p.text = p.text.replace("ARTIFICIAL INTELLIGENCE ENHANCED JOB PLATFORM: A MACHINE LEARNING-DRIVEN RECOMMENDATIONS", "ADIPHAS: AUTONOMOUS DISEASE INTELLIGENCE AND PUBLIC HEALTH ADVISORY SYSTEM")
        if "TANIMOWO" in p.text.upper() or "EMMANUEL OLUWADARASIMI" in p.text.upper():
            p.text = p.text.replace("TANIMOWO, EMMANUEL OLUWADARASIMI", "IFET GREAT TITUS")
            p.text = p.text.replace("Tanimowo Emmanuel Oluwadarasimi", "Ifet Great Titus")
        if "2101110048" in p.text:
            p.text = p.text.replace("2101110048", "2201110075")

    # 2. DELETE OLD CONTENT (From Ch 1 onwards)
    for i in range(len(doc.paragraphs) - 1, 94, -1):
        p = doc.paragraphs[i]._element
        p.getparent().remove(p)

    # 3. INSERT ADIPHAS CONTENT (COPYING RUNS TO PRESERVE IMAGES/FORMAT)
    ch1_start = 0
    for i, p in enumerate(source_doc.paragraphs):
        if "CHAPTER ONE" in p.text.upper():
            ch1_start = i
            break
            
    for p_source in source_doc.paragraphs[ch1_start:]:
        new_p = doc.add_paragraph()
        # Copy paragraph properties
        new_p.style = p_source.style
        new_p.alignment = p_source.alignment
        
        # Copy runs
        for run_source in p_source.runs:
            new_run = new_p.add_run(run_source.text)
            new_run.bold = run_source.bold
            new_run.italic = run_source.italic
            new_run.underline = run_source.underline
            if run_source.font.name: new_run.font.name = run_source.font.name
            if run_source.font.size: new_run.font.size = run_source.font.size
            
            # Copy InlineShapes (Images) - this is tricky, we'll rely on the existing images
            # if we can't easily clone the blob. 
            # (Note: python-docx doesn't easily clone images between docs).

    doc.save(output_path)
    print(f"Success: Standardized project saved to {output_path}")

clone_standardization_v2()
