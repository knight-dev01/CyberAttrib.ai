
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import os

def add_toc(paragraph):
    # Inserts a standard Word TOC field
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._element.append(fldChar)

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    run._element.append(instrText)

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    run._element.append(fldChar2)

    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._element.append(fldChar3)

def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.xpath('w:tblPr')[0]
    tblBorders = OxmlElement('w:tblBorders')
    for border in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = OxmlElement(f'w:{border}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), '000000')
        tblBorders.append(b)
    tblPr.append(tblBorders)

def final_perfection():
    doc_path = r'c:\Users\Adegoke\Desktop\New folder\docs\ADIPHAS\ADIPHAS_v6_Final_Titus_Submission.docx'
    output_path = r'c:\Users\Adegoke\Desktop\New folder\docs\ADIPHAS\ADIPHAS_Final_Submission_Perfected.docx'
    md_path = r'c:\Users\Adegoke\Desktop\New folder\docs\ADIPHAS\PROJECT_DOCUMENTATION.md'
    
    doc = docx.Document(doc_path)
    with open(md_path, 'r', encoding='utf-8') as f:
        md_content = f.read()

    # 1. ADD TABLE OF CONTENTS (After Acknowledgement)
    for i, p in enumerate(doc.paragraphs):
        if "ACKNOWLEDGEMENT" in p.text.upper():
            # Add a page break and TOC
            p.add_run().add_break(docx.enum.text.WD_BREAK.PAGE)
            toc_title = doc.add_paragraph("TABLE OF CONTENTS")
            toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            toc_title.runs[0].bold = True
            
            toc_p = doc.add_paragraph()
            add_toc(toc_p)
            
            # Move TOC elements to after Acknowledgement
            toc_title_elem = toc_title._element
            toc_p_elem = toc_p._element
            p._element.addnext(toc_title_elem)
            toc_title_elem.addnext(toc_p_elem)
            break

    # 2. ADD COMPARISON ANALYSIS (End of Chapter 4)
    # Extract data
    table_data = [
        ["Capability Dimension", "ADIPHAS (Our System)", "SORMAS (NCDC Standard)", "HealthMap (Global)", "ProMED-mail"],
        ["Data Acquisition Model", "Autonomous EBS: Silently scrapes 20+ sources.", "Indicator-Based (IBS): Relies on clinical case reporting.", "Global EBS: Large scale, lacks local Nigerian depth.", "Manual EBS: Crowdsourced reports via email."],
        ["Latency (Lead-Time)", "9-11 Days Lead Time (Proactive).", "Reactive (Days to weeks after clinical presentation).", "Variable (High latency for local outbreaks).", "High Latency (Moderated reports)."],
        ["Resilience", "25-model failover & HTTP API bypass.", "Human-Dependent (Manual data entry).", "Centralized Cloud (Single point of failure).", "Human-Dependent (Manual moderation)."],
        ["Hyper-Local Personalization", "Browser-native GPS + LGA Risk Scoring.", "National/State focus only.", "Global/City overview only.", "Technical/Academic summaries."]
    ]
    
    for i, p in enumerate(doc.paragraphs):
        if "CHAPTER FIVE" in p.text.upper():
            # Insert before Chapter 5
            sect_p = doc.add_paragraph("4.5 Comprehensive Comparative Analysis")
            sect_p.style = 'Heading 2'
            sect_p.runs[0].bold = True
            
            p._element.addprevious(sect_p._element)
            
            # Add Table
            table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
            set_table_borders(table)
            for r_idx, row_data in enumerate(table_data):
                for c_idx, cell_text in enumerate(row_data):
                    table.cell(r_idx, c_idx).text = cell_text
            p._element.addprevious(table._element)
            
            # Add Architectural Superiority points
            superiority_text = [
                "1. Zero-Touch Surveillance: Unlike legacy systems, ADIPHAS operates strictly in the background, freeing clinicians for patient care.",
                "2. Defeating the 'Cold Start': ADIPHAS guarantees 24/7 vigilance without human fatigue.",
                "3. Hyper-Personalization: Hyper-tailored advice filtered by biological markers (Genotype/Blood Group).",
                "4. Resilience Engineering: 25-model exponential backoff matrix ensures high availability."
            ]
            for text in superiority_text:
                para = doc.add_paragraph(text)
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p._element.addprevious(para._element)
            break

    # 3. SCREENSHOTS SPACE
    for p in doc.paragraphs:
        if "4.4 Screenshots of the Developed System" in p.text:
            # Add empty paragraphs for space
            for _ in range(3):
                empty_p = doc.add_paragraph("\n[SPACE FOR SCREENSHOT]\n")
                empty_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p._element.addnext(empty_p._element)

    # 4. FINAL AUDIT & FORMATTING
    for p in doc.paragraphs:
        if p._element.xpath('.//w:drawing'): continue
        if "REFERENCES" in p.text.upper(): break
        
        # Consistent spacing and font
        p.paragraph_format.line_spacing = 1.5
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in p.runs:
            if not run.font.name: run.font.name = 'Times New Roman'
            if not run.font.size: run.font.size = Pt(12)

    doc.save(output_path)
    print(f"Success: Final perfected document saved to {output_path}")

    # CLEANUP
    files_to_delete = [
        r'c:\Users\Adegoke\Desktop\New folder\update_adiphas_xml.py',
        r'c:\Users\Adegoke\Desktop\New folder\update_adiphas_surgical_v2.py',
        r'c:\Users\Adegoke\Desktop\New folder\update_adiphas_surgical_v3.py',
        r'c:\Users\Adegoke\Desktop\New folder\sync_adiphas_metrics.py',
        r'c:\Users\Adegoke\Desktop\New folder\sync_adiphas_metrics_v2.py',
        r'c:\Users\Adegoke\Desktop\New folder\master_refactor_adiphas.py',
        r'c:\Users\Adegoke\Desktop\New folder\finalize_adiphas_submission.py',
        r'c:\Users\Adegoke\Desktop\New folder\surgical_final_v1.py',
        r'c:\Users\Adegoke\Desktop\New folder\docs\ADIPHAS\ADIPHAS_v6_Updated_Professional.docx',
        r'c:\Users\Adegoke\Desktop\New folder\docs\ADIPHAS\ADIPHAS_v6_Final_Perfected.docx',
        r'c:\Users\Adegoke\Desktop\New folder\docs\ADIPHAS\ADIPHAS_v6_Final_Metrics_Sync.docx'
    ]
    for f_path in files_to_delete:
        if os.path.exists(f_path):
            try: os.remove(f_path)
            except: pass

final_perfection()
